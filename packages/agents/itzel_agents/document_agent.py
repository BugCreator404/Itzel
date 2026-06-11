"""
DocumentAgent — lectura y escritura de documentos para Itzel.

Extrae texto de PDF, DOCX y XLSX, crea documentos Word y resume documentos
usando el modelo Itzel-1B (o un resumen extractivo de respaldo si el modelo
no está conectado).

Dependencias opcionales (extra [documents]):
    pdfminer.six   → read_pdf
    python-docx    → read_docx / write_docx
    openpyxl       → read_xlsx

Degradación elegante:
    Si falta una dependencia, la tool devuelve un error legible con la
    instrucción de instalación, sin romper el agente.

Seguridad:
    write_docx confirma solo si el archivo de destino ya existe (sobrescritura).

English summary:
    Document agent: extract text from PDF/DOCX/XLSX, write DOCX, and summarize
    via Itzel-1B (with an extractive fallback). Optional deps degrade cleanly.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Callable, Optional

from itzel_core.tools import Tool, tool

from .base_agent import BaseAgent

# ─── límites ──────────────────────────────────────────────────────────────────

_MAX_DOC_CHARS     = 200_000   # tope de texto extraído (evita explotar memoria)
_SUMMARY_SENTENCES = 5         # oraciones del resumen extractivo de respaldo
_XLSX_MAX_ROWS     = 1_000     # filas máximas a leer por hoja


def _resolve(path: str, *, must_exist: bool = False) -> Path:
    return BaseAgent._resolve_path(path, must_exist=must_exist)


def _path_exists(args: dict[str, Any]) -> bool:
    """Predicado de confirmación para sobrescritura de 'path'."""
    p = args.get("path")
    if not p:
        return False
    try:
        return _resolve(str(p)).exists()
    except Exception:
        return True


def _clip(text: str) -> str:
    """Recorta el texto extraído al tope de seguridad."""
    if len(text) > _MAX_DOC_CHARS:
        return text[:_MAX_DOC_CHARS] + "\n…[texto truncado]"
    return text


# ─── tools: lectura ───────────────────────────────────────────────────────────

@tool(
    "read_pdf",
    "Extrae el texto de un archivo PDF.",
    timeout=60,
    param_docs={"path": "Ruta del archivo PDF."},
)
def read_pdf(path: str) -> str:
    """Extrae texto de un PDF usando pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text
    except ImportError:
        raise RuntimeError(
            "pdfminer.six no está instalado. "
            "Instala con: pip install -e packages/agents[documents]"
        )
    p = _resolve(path, must_exist=True)
    text = extract_text(str(p)) or ""
    return _clip(text)


@tool(
    "read_docx",
    "Extrae el texto de un documento Word (.docx).",
    param_docs={"path": "Ruta del archivo .docx."},
)
def read_docx(path: str) -> str:
    """Extrae texto de un .docx (párrafos y tablas) con python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx no está instalado. "
            "Instala con: pip install -e packages/agents[documents]"
        )
    p = _resolve(path, must_exist=True)
    doc = Document(str(p))

    parts: list[str] = [para.text for para in doc.paragraphs if para.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    return _clip("\n".join(parts))


@tool(
    "read_xlsx",
    "Extrae los datos de una hoja de cálculo Excel (.xlsx).",
    param_docs={"path": "Ruta del archivo .xlsx."},
)
def read_xlsx(path: str) -> dict[str, Any]:
    """Lee un .xlsx con openpyxl. Devuelve filas por hoja (con límite)."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise RuntimeError(
            "openpyxl no está instalado. "
            "Instala con: pip install -e packages/agents[documents]"
        )
    p = _resolve(path, must_exist=True)
    wb = load_workbook(str(p), read_only=True, data_only=True)

    sheets: dict[str, list[list[Any]]] = {}
    for ws in wb.worksheets:
        rows: list[list[Any]] = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= _XLSX_MAX_ROWS:
                break
            rows.append([("" if c is None else c) for c in row])
        sheets[ws.title] = rows
    wb.close()

    return {"sheets": sheets, "sheet_names": list(sheets.keys())}


# ─── tools: escritura ─────────────────────────────────────────────────────────

@tool(
    "write_docx",
    "Crea un documento Word (.docx) con el contenido dado. Confirma si ya existe.",
    confirm_when=_path_exists,
    param_docs={
        "path":    "Ruta del .docx a crear.",
        "content": "Texto del documento (cada línea es un párrafo).",
    },
)
def write_docx(path: str, content: str) -> str:
    """Crea un .docx con python-docx. Cada salto de línea es un párrafo."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx no está instalado. "
            "Instala con: pip install -e packages/agents[documents]"
        )
    p = _resolve(path)
    p.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(str(p))
    return f"Documento creado: {p}"


# ─── resumen extractivo de respaldo ───────────────────────────────────────────

def _extractive_summary(text: str, n_sentences: int = _SUMMARY_SENTENCES) -> str:
    """
    Resumen de respaldo sin LLM: toma las primeras N oraciones no vacías.
    Se usa cuando el modelo Itzel-1B no está conectado.
    """
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    picked = [s.strip() for s in sentences if s.strip()][:n_sentences]
    summary = " ".join(picked)
    return summary or "[documento vacío o sin texto extraíble]"


def _extract_text_any(path: str) -> str:
    """Detecta el tipo de documento por extensión y extrae su texto."""
    p = _resolve(path, must_exist=True)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(str(p))
    if suffix == ".docx":
        return read_docx(str(p))
    if suffix in (".txt", ".md", ".csv"):
        return _clip(p.read_text(encoding="utf-8", errors="replace"))
    if suffix == ".xlsx":
        data = read_xlsx(str(p))
        lines = []
        for sheet, rows in data["sheets"].items():
            lines.append(f"# {sheet}")
            for row in rows:
                lines.append("\t".join(str(c) for c in row))
        return _clip("\n".join(lines))
    raise ValueError(f"Tipo de documento no soportado para resumen: '{suffix}'")


# ─── el agente ────────────────────────────────────────────────────────────────

class DocumentAgent(BaseAgent):
    """
    Agente de documentos. 5 tools: read_pdf, read_docx, read_xlsx, write_docx,
    summarize.

    summarize usa el `summarizer` inyectado (típicamente un wrapper de Itzel-1B);
    si no se provee, cae a un resumen extractivo simple.
    """

    name = "document"

    def __init__(
        self,
        *,
        summarizer: Optional[Callable[[str], str]] = None,
        **kwargs: Any,
    ) -> None:
        self._summarizer = summarizer
        super().__init__(**kwargs)

    def _register_tools(self) -> None:
        self.registry.register_all([read_pdf, read_docx, read_xlsx, write_docx])
        self.registry.register(self._make_summarize_tool())

    def _make_summarize_tool(self) -> Tool:
        """Construye la tool summarize capturando el summarizer del agente."""
        summarizer = self._summarizer

        @tool(
            "summarize",
            "Resume un documento (PDF, DOCX, XLSX, TXT) usando Itzel-1B.",
            timeout=120,
            param_docs={"path": "Ruta del documento a resumir."},
        )
        def summarize(path: str) -> str:
            text = _extract_text_any(path)
            if not text.strip():
                return "[documento vacío o sin texto extraíble]"
            if summarizer is None:
                # Sin LLM conectado: resumen extractivo de respaldo.
                return _extractive_summary(text)
            return summarizer(text)

        return summarize
